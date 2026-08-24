import sys
import json
import os
from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph


def read_docx(path, **_):
    doc = Document(path)
    paragraphs = [
        {"text": p.text, "style": p.style.name if p.style else None}
        for p in doc.paragraphs
    ]
    tables = []
    for table in doc.tables:
        tables.append([[cell.text for cell in row.cells] for row in table.rows])
    return {"paragraphs": paragraphs, "tables": tables}


def create_docx(path, paragraphs=None, overwrite=False, **_):
    if os.path.exists(path) and not overwrite:
        raise RuntimeError(f"檔案已存在，未帶 overwrite:true 不會覆蓋: {path}")
    doc = Document()
    for item in paragraphs or []:
        text = item.get("text", "")
        heading_level = item.get("heading_level")
        if heading_level is not None:
            doc.add_heading(text, level=heading_level)
        else:
            p = doc.add_paragraph(text)
            style = item.get("style")
            if style:
                p.style = doc.styles[style]
    doc.save(path)
    return {"path": path}


def append_paragraph(path, paragraphs=None, **_):
    if not os.path.exists(path):
        raise RuntimeError(f"檔案不存在，請先用 create_docx 建立: {path}")
    doc = Document(path)
    for item in paragraphs or []:
        text = item.get("text", "")
        heading_level = item.get("heading_level")
        if heading_level is not None:
            doc.add_heading(text, level=heading_level)
        else:
            p = doc.add_paragraph(text)
            style = item.get("style")
            if style:
                p.style = doc.styles[style]
    doc.save(path)
    return {"path": path, "appended": len(paragraphs or [])}


def replace_text(path, find, replace, **_):
    if not os.path.exists(path):
        raise RuntimeError(f"檔案不存在: {path}")
    doc = Document(path)
    count = 0

    def replace_in_paragraph(p):
        nonlocal count
        full_text = "".join(run.text for run in p.runs)
        if find not in full_text:
            return
        new_text = full_text.replace(find, replace)
        count += full_text.count(find)
        # 找出的取代結果整段塞回第一個 run、清空其餘 run——docx 的文字常被拆成多個 run（不同格式片段），
        # 逐個 run 比對容易漏掉跨 run 的字串，這裡犧牲原有的 run 級格式（保留段落層級格式）換取比對正確性。
        if p.runs:
            p.runs[0].text = new_text
            for run in p.runs[1:]:
                run.text = ""

    for p in doc.paragraphs:
        replace_in_paragraph(p)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    doc.save(path)
    return {"path": path, "replacements": count}


def insert_table(path, rows, style=None, **_):
    if not os.path.exists(path):
        raise RuntimeError(f"檔案不存在，請先用 create_docx 建立: {path}")
    doc = Document(path)
    if not rows:
        raise RuntimeError("rows 不能是空陣列")
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if style:
        table.style = style
    for r, row_values in enumerate(rows):
        for c, value in enumerate(row_values):
            table.cell(r, c).text = str(value)
    doc.save(path)
    return {"path": path, "rows": len(rows), "cols": len(rows[0])}


def _iter_block_items(doc):
    """依文件實際順序（含表格內、跳過表格外）依序 yield 段落與表格物件，比分開列舉 doc.paragraphs / doc.tables 更準確反映閱讀順序，讓標題/前後文關聯正確。"""
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            yield Paragraph(child, doc)
        elif child.tag == qn("w:tbl"):
            yield Table(child, doc)


def _find_blips(paragraph):
    return paragraph._element.findall(".//" + qn("a:blip"))


def extract_images(path, output_dir, **_):
    if not os.path.exists(path):
        raise RuntimeError(f"檔案不存在: {path}")
    os.makedirs(output_dir, exist_ok=True)
    doc = Document(path)

    state = {"heading": None, "last_text": None}
    images = []
    pending_after = []

    def save_blip(blip, extra):
        r_embed = blip.get(qn("r:embed"))
        if not r_embed:
            return None
        rel = doc.part.rels.get(r_embed)
        if rel is None or "image" not in rel.reltype:
            return None
        seq = len(images) + 1
        ext = os.path.splitext(rel.target_ref)[1] or ".png"
        out_path = os.path.join(output_dir, f"image_{seq:03d}{ext}")
        with open(out_path, "wb") as f:
            f.write(rel.target_part.blob)
        record = {
            "file": out_path,
            "section_heading": state["heading"],
            "context_before": state["last_text"],
            "context_after": None,
            **extra,
        }
        images.append(record)
        return record

    def visit_paragraph(p, extra):
        text = p.text.strip()
        for blip in _find_blips(p):
            record = save_blip(blip, extra)
            if record is not None:
                pending_after.append(record)
        if text:
            if p.style and p.style.name and p.style.name.startswith("Heading"):
                state["heading"] = text
            for record in pending_after:
                record["context_after"] = text
            pending_after.clear()
            state["last_text"] = text

    for block in _iter_block_items(doc):
        if isinstance(block, Table):
            for r, row in enumerate(block.rows):
                for c, cell in enumerate(row.cells):
                    for p in cell.paragraphs:
                        visit_paragraph(p, {"in_table": True, "table_row": r, "table_col": c})
        else:
            visit_paragraph(block, {"in_table": False})

    return {"images": images, "count": len(images)}


ACTIONS = {
    "read": read_docx,
    "create": create_docx,
    "append_paragraph": append_paragraph,
    "replace_text": replace_text,
    "insert_table": insert_table,
    "extract_images": extract_images,
}


def main():
    payload = json.loads(sys.stdin.buffer.read().decode("utf-8"))
    action = payload.get("action")
    if action not in ACTIONS:
        raise RuntimeError(f"未知的 action: {action}（可用: {', '.join(ACTIONS)}）")
    result = ACTIONS[action](**{k: v for k, v in payload.items() if k != "action"})
    print(json.dumps({"success": True, **result}, ensure_ascii=False))


if __name__ == "__main__":
    sys.stdout.reconfigure(encoding="utf-8")
    try:
        main()
    except Exception as e:
        print(json.dumps({"success": False, "message": str(e)}, ensure_ascii=False))
        sys.exit(1)
